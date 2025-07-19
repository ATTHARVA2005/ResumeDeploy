# tests/test_resume_parser.py

import pytest
import os
from pathlib import Path
import shutil

# Import the ResumeParser class from the backend module
# Adjust the import path if your project structure changes
from backend.resume_parser import ResumeParser
from backend.utils import save_upload_file, delete_file # For creating/cleaning up test files

# Define a temporary directory for test files
TEST_FILES_DIR = "test_temp_files"

@pytest.fixture(scope="module", autouse=True)
def setup_and_teardown_test_files():
    """
    Fixture to create and clean up a temporary directory for test files.
    This runs once before all tests in this module and once after.
    """
    # Setup: Create a temporary directory
    os.makedirs(TEST_FILES_DIR, exist_ok=True)
    print(f"\nCreated test directory: {TEST_FILES_DIR}")
    
    # Yield control to the tests
    yield
    
    # Teardown: Clean up the temporary directory
    if os.path.exists(TEST_FILES_DIR):
        shutil.rmtree(TEST_FILES_DIR)
        print(f"Cleaned up test directory: {TEST_FILES_DIR}")

@pytest.fixture
def resume_parser_instance():
    """Fixture to provide a ResumeParser instance for tests."""
    return ResumeParser()

@pytest.mark.asyncio # Mark async tests for pytest-asyncio
async def test_extract_text_pdf(resume_parser_instance):
    """Test PDF text extraction with a dummy PDF file."""
    # Create a dummy PDF file for testing
    # In a real scenario, you'd have a small, actual PDF file here.
    # For now, we'll simulate a file path.
    dummy_pdf_content = "%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj 2 0 obj<</Type/Pages/Count 1/Kids[3 0 R]>>endobj 3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Contents 4 0 R/Parent 2 0 R>>endobj 4 0 obj<</Length 44>>stream\nBT /F1 24 Tf 100 700 Td (Hello, PDF world!) Tj ET\nendstream\nendobj\nxref\n0 5\n0000000000 65535 f\n0000000009 00000 n\n0000000055 00000 n\n0000000109 00000 n\n0000000201 00000 n\ntrailer<</Size 5/Root 1 0 R>>startxref\n250\n%%EOF"
    pdf_file_path = os.path.join(TEST_FILES_DIR, "dummy_resume.pdf")
    with open(pdf_file_path, "wb") as f:
        f.write(dummy_pdf_content.encode('latin-1')) # PDF content often needs specific encoding

    extracted_text = resume_parser_instance.extract_text(pdf_file_path)
    assert "Hello, PDF world!" in extracted_text
    assert "pdf" in extracted_text.lower() # Check for some common PDF related text if not exactly 'Hello...'

    # Clean up the dummy file
    delete_file(pdf_file_path)

@pytest.mark.asyncio
async def test_extract_text_docx(resume_parser_instance):
    """Test DOCX text extraction with a dummy DOCX file."""
    # python-docx creates valid DOCX files
    from docx import Document
    doc = Document()
    doc.add_paragraph("This is a test DOCX resume.")
    doc.add_paragraph("It contains some skills like Python and Java.")
    docx_file_path = os.path.join(TEST_FILES_DIR, "dummy_resume.docx")
    doc.save(docx_file_path)

    extracted_text = resume_parser_instance.extract_text(docx_file_path)
    assert "This is a test DOCX resume." in extracted_text
    assert "Python" in extracted_text
    assert "Java" in extracted_text

    delete_file(docx_file_path)

@pytest.mark.asyncio
async def test_extract_text_txt(resume_parser_instance):
    """Test TXT text extraction."""
    txt_content = "This is a plain text resume.\nSkills: C++, SQL, AWS."
    txt_file_path = os.path.join(TEST_FILES_DIR, "dummy_resume.txt")
    with open(txt_file_path, "w", encoding="utf-8") as f:
        f.write(txt_content)

    extracted_text = resume_parser_instance.extract_text(txt_file_path)
    assert "This is a plain text resume." in extracted_text
    assert "C++" in extracted_text
    assert "SQL" in extracted_text

    delete_file(txt_file_path)

def test_clean_text(resume_parser_instance):
    """Test text cleaning functionality."""
    dirty_text = "  Hello   World! \n\nThis is a test.  \tWith\xa0special\u2022chars."
    cleaned_text = resume_parser_instance._clean_text(dirty_text)
    assert cleaned_text == "Hello World! This is a test. With special chars."

    dirty_text_2 = "Text with/slashes:and:colons, and+plus#hashes"
    cleaned_text_2 = resume_parser_instance._clean_text(dirty_text_2)
    assert cleaned_text_2 == "Text with/slashes:and:colons, and+plus#hashes"


def test_extract_contact_info(resume_parser_instance):
    """Test extraction of contact information."""
    text = """
    Contact: john.doe@example.com | +1 (123) 456-7890
    LinkedIn: linkedin.com/in/john-doe-profile
    GitHub: github.com/johndoe
    Another Email: jane.smith@domain.co.uk
    Phone: 987-654-3210
    """
    contact_info = resume_parser_instance.extract_contact_info(text)

    assert "john.doe@example.com" in contact_info['emails']
    assert "jane.smith@domain.co.uk" in contact_info['emails']
    assert len(contact_info['emails']) == 2

    assert "+1 (123) 456-7890" in contact_info['phones']
    assert "987-654-3210" in contact_info['phones']
    assert len(contact_info['phones']) == 2

    assert "linkedin.com/in/john-doe-profile" in contact_info['linkedin'][0] # Check for substring due to potential protocol variations
    assert "github.com/johndoe" in contact_info['github'][0]

def test_extract_sections(resume_parser_instance):
    """Test extraction of resume sections."""
    text = """
    John Doe
    Software Engineer

    SUMMARY
    Highly skilled engineer.

    EDUCATION
    University of XYZ - BS in CS

    WORK EXPERIENCE
    Company A - Software Developer
    Developed features using Python.

    PROJECTS
    My Awesome Project - Built with React.

    SKILLS
    Python, Java, AWS, Docker.
    """
    sections = resume_parser_instance.extract_sections(text)

    assert "personal_info" in sections
    assert "John Doe" in sections['personal_info'][0]
    assert "Software Engineer" in sections['personal_info'][0]

    assert "summary_objective" in sections
    assert "Highly skilled engineer." in sections['summary_objective'][0]

    assert "education" in sections
    assert "University of XYZ - BS in CS" in sections['education'][0]

    assert "experience" in sections
    assert "Company A - Software Developer" in sections['experience'][0]
    assert "Developed features using Python." in sections['experience'][0]

    assert "projects" in sections
    assert "My Awesome Project - Built with React." in sections['projects'][0]

    assert "skills" in sections
    assert "Python, Java, AWS, Docker." in sections['skills'][0]

    assert "other" not in sections # Should not have 'other' if all sections are found
