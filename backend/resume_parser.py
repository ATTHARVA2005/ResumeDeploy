# backend/resume_parser.py

import os
import re
import json
from datetime import datetime
from typing import List, Dict, Any, Optional
from pathlib import Path

# Text extraction libraries
from pdfminer.high_level import extract_text as pdf_extract_text
from docx import Document

# New: Import Gemini API client
import google.generativeai as genai

# New: Configure Gemini API key from environment variable
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

class ResumeParser: # Keeping the class name as ResumeParser for now, can be renamed to GeminiExtractor if preferred
    def __init__(self):
        pass
    
    def extract_text(self, file_path: str) -> str:
        """Extracts plain text from various resume file formats."""
        file_extension = Path(file_path).suffix.lower()
        extracted_text = ""

        try:
            if file_extension == '.pdf':
                extracted_text = self._extract_from_pdf(file_path)
            elif file_extension == '.docx':
                extracted_text = self._extract_from_docx(file_path)
            elif file_extension == '.txt':
                extracted_text = self._extract_from_txt(file_path)
            else:
                print(f"Unsupported extension '{file_extension}' for direct text extraction.")
                extracted_text = "" 
        except Exception as e:
            print(f"Error during text extraction from {file_path}: {e}")
            extracted_text = ""

        return self._clean_text(extracted_text)

    def _extract_from_pdf(self, file_path: str) -> str:
        """Extracts text from PDF files using pdfminer.six."""
        try:
            text = pdf_extract_text(file_path)
            return text
        except Exception as e:
            print(f"PDF extraction error for {file_path}: {e}")
            return ""
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extracts text from DOCX files using python-docx."""
        try:
            doc = Document(file_path)
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text)
            return '\n'.join(full_text)
        except Exception as e:
            print(f"DOCX extraction error for {file_path}: {e}")
            return ""
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extracts text from TXT files."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
            return text
        except Exception as e:
            print(f"TXT extraction error for {file_path}: {e}")
            return ""
    
    def _clean_text(self, text: str) -> str:
        """Cleans and normalizes extracted text for better NLP processing."""
        if not text:
            return ""
        
        text = text.replace('\xa0', ' ').replace('\u2022', ' ').replace('\u2013', '-')
        text = re.sub(r'\$[0-9]+\^{.+?}\$', '', text) 
        text = re.sub(r'\"', '', text)
        text = re.sub(r',,+', ',', text)
        text = re.sub(r'\s{2,}', ' ', text)
        text = re.sub(r'\n{2,}', '\n', text)
        text = re.sub(r'^\s*,\s*', '', text, flags=re.MULTILINE)
        text = re.sub(r',\s*$', '', text, flags=re.MULTILINE)
        
        text = re.sub(r'\n+', '\n', text) 
        text = '\n'.join([line.strip() for line in text.split('\n')])
        
        return text.strip()

    async def parse_text_with_gemini(self, raw_text: str) -> Dict[str, Any]:
        """
        Sends extracted raw text of a resume to Gemini for structured data extraction.
        
        Args:
            raw_text (str): The cleaned text extracted from a resume.
            
        Returns:
            Dict[str, Any]: A dictionary containing structured resume data.
        """
        model = genai.GenerativeModel('gemma-3-12b-it')

        json_schema = """
        {
          "full_name": "String (e.g., John Doe)",
          "email": "String (e.g., john.doe@example.com)",
          "phone": "String (e.g., +1-555-123-4567)",
          "linkedin_url": "String (e.g., https://www.linkedin.com/in/johndoe)",
          "github_url": "String (e.g., https://github.com/johndoe)",
          "total_years_experience": "Integer (e.g., 5, calculated from work experience dates. Return 0 if no experience found)",
          "highest_education_level": "String (e.g., Master's, PhD, Bachelor's, Diploma, High School. Return 'None' if not found)",
          "major": "String (e.g., Computer Science, Electrical Engineering. Return 'None' if not found)",
          "extracted_skills": [
            "Skill 1 (e.g., Python)",
            "Skill 2 (e.g., Data Analysis)",
            "Skill 3 (e.g., SQL)"
          ],
          "experience": [
            {
              "title": "String (e.g., Senior Software Engineer)",
              "company": "String (e.g., Google)",
              "start_date": "String (e.g., Jan 2020, 2020. If only year, just year)",
              "end_date": "String (e.g., Present, Dec 2023, 2023. If only year, just year)",
              "description": "String (concise bullet points/paragraph of responsibilities and achievements. Keep key action verbs and measurable results.)"
            }
          ]
        }
        """

        prompt = f"""
        Analyze the following resume text and extract the specified information into a JSON object.
        Ensure the JSON strictly adheres to the provided schema.
        
        If a field is not found, provide an appropriate default value as indicated in the schema (e.g., empty string, empty list, 0, or "None").
        For 'total_years_experience', calculate it based on the 'start_date' and 'end_date' of all experience entries, assuming 'Present' means current date.
        
        Resume Text:
        ---
        {raw_text}
        ---
        
        JSON Schema to follow:
        {json_schema}
        
        Provide only the JSON object.
        """
        
        try:
            response = await model.generate_content_async(prompt)
            
            response_text = response.text.strip()
            if response_text.startswith("```json") and response_text.endswith("```"):
                json_string = response_text[len("```json"):-len("```")].strip()
            else:
                json_string = response_text

            parsed_data = json.loads(json_string)
            return parsed_data
            
        except Exception as e:
            print(f"Error calling Gemini API or parsing resume response: {e}")
            return {
                "full_name": "",
                "email": "",
                "phone": "",
                "linkedin_url": "",
                "github_url": "",
                "total_years_experience": 0,
                "highest_education_level": "None",
                "major": "None",
                "extracted_skills": [],
                "experience": [],
            }

    async def parse_job_description_with_gemini(self, raw_text: str) -> Dict[str, Any]:
        """
        Sends raw job description text to Gemini for structured requirements extraction.
        
        Args:
            raw_text (str): The full text of the job description.
            
        Returns:
            Dict[str, Any]: A dictionary containing structured job requirements data.
        """
        model = genai.GenerativeModel('gemma-3-12b-it')

        json_schema = """
        {          
          "description": "String (The full job description text provided as input)",
          "required_experience_years": "Integer (e.g., 5. Infer from text like '5+ years experience', 'minimum 3 years', 'entry-level'. Return 0 if not specified, 0 for entry-level.)",
          "required_skills": [
            "Skill 1 (e.g., Python)",
            "Skill 2 (e.g., Data Analysis)",
            "Skill 3 (e.g., SQL)"
          ],
          "required_certifications": [
            "Certification 1 (e.g., AWS Certified Solutions Architect)",
            "Certification 2 (e.g., PMP)"
          ],
          "required_education_level": "String (e.g., Bachelor's, Master's, PhD, High School Diploma. Infer from text like 'Bachelor's degree required', 'Master's preferred'. Return 'None' if not specified)",
          "required_major": "String (e.g., Computer Science, Electrical Engineering. Infer from text like 'degree in CS', 'background in engineering'. Return 'None' if not specified)"
        }
        """

        prompt = f"""
        Analyze the following job description text and extract the specified requirements into a JSON object.
        Ensure the JSON strictly adheres to the provided schema.
        
        If a field is not found or not explicitly mentioned, provide an appropriate default value as indicated in the schema (e.g., empty string, empty list, 0, or "None").
        Infer 'required_experience_years' carefully from phrases like 'X years of experience', 'X+ years', 'entry-level', 'junior', 'senior'. If 'entry-level' or similar, return 0.
        For 'required_skills' and 'required_certifications', provide a list of distinct, relevant items.
        
        Job Description Text:
        ---
        {raw_text}
        ---
        
        JSON Schema to follow:
        {json_schema}
        
        Provide only the JSON object.
        """

        try:
            response = await model.generate_content_async(prompt)
            
            response_text = response.text.strip()
            if response_text.startswith("```json") and response_text.endswith("```"):
                json_string = response_text[len("```json"):-len("```")].strip()
            else:
                json_string = response_text

            parsed_data = json.loads(json_string)
            
            # Ensure required_experience_years is an integer
            if 'required_experience_years' in parsed_data:
                try:
                    parsed_data['required_experience_years'] = int(parsed_data['required_experience_years'])
                except (ValueError, TypeError):
                    parsed_data['required_experience_years'] = 0 # Default to 0 if conversion fails

            return parsed_data
            
        except Exception as e:
            print(f"Error calling Gemini API or parsing job description response: {e}")
            # Return a default empty structure for fallback
            return {
                "title": "",
                "company": "",
                "description": raw_text, # Keep the original description
                "required_experience_years": 0,
                "required_skills": [],
                "required_certifications": [],
                "required_education_level": "None",
                "required_major": "None"
            }